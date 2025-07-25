from docx import Document
import pandas as pd
import os

data_path = "data"
os.makedirs(data_path, exist_ok=True)

# Create DOCX file
doc = Document()
doc.add_heading("Quarterly KPI Report", 0)
doc.add_paragraph("- Customer Retention Rate: 82%")
doc.add_paragraph("- Net Promoter Score (NPS): 45")
doc.add_paragraph("- Monthly Active Users: 12,000")
doc.save(os.path.join(data_path, "report.docx"))

# Create CSV file
df = pd.DataFrame({
    "KPI": ["Customer Retention Rate", "Net Promoter Score", "Monthly Active Users"],
    "Q1": [82, 45, 12000],
    "Q2": [80, 47, 12500]
})
df.to_csv(os.path.join(data_path, "metrics.csv"), index=False)

# Create TXT file
with open(os.path.join(data_path, "test.txt"), "w") as f:
    f.write("Customer Retention Rate was 82% in Q1.\n")
    f.write("Net Promoter Score improved to 45.\n")
    f.write("Monthly Active Users reached 12,000.\n")

print("✅ Files created successfully in 'data/'")
