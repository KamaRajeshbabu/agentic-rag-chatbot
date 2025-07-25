from docx import Document

doc = Document()
doc.add_heading("Sample Report", 0)
doc.add_paragraph("These are some KPIs tracked in Q1:")
doc.add_paragraph("- Revenue Growth\n- Customer Retention Rate\n- Net Profit Margin")
doc.save("data/sample_files/docx_sample.docx")

